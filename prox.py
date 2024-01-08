from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture('itachi.png', width=Inches(2.0))

# name phone number and email address
name = input('what is your name ?')
speak('hello' + name + 'how are you today')

speak('Please enter your phone number to continue to next step')
phone_number = input('what is your phone number ?')

speak('Please enter your email address')
email = input('what is your email id ?')


document.add_paragraph(name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About me')
speak('Describe yourself' + name +' ')
about_me = input('Describe  yourself')
document.add_paragraph(about_me)

# organization experience
document.add_heading('University')
p = document.add_paragraph()

speak('enter your college name')
college = input('Enter college Name')

speak('enter your couse details')
course = input('Which course are you studying')

speak('enter your couse duration')
cstart = input('Course starting year')
cend = input('Course ending year')

p.add_run(college + ' ').bold = True
p.add_run(course + ' , ').bold = True
p.add_run(cstart + '-' + cend + '\n').italic = True

speak('describe your status at your college')
experience_details = input('Describe your current status at ' + college + '\n')
p.add_run(experience_details)

# skills
document.add_heading('Skills')

speak('enter your skill')
skills = input('Enter Skill')
p = document.add_paragraph(skills)

speak('if any other skills you have please type yes otherwise type no')
while True:
    has_more_skills = input('Do you have more skills? Yes or No ')
    if has_more_skills.lower() == 'yes' :
        skills = input('Enter skill')
        p = document.add_paragraph(skills)
        p.style = 'List Bullet'
    else:
        break

# footer 
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using project X software by Tony Sensei"

document.save('cv.docx')